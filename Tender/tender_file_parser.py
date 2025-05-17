"""
Tender File Parser Utility

This script provides functions to parse tender information from CSV, Excel, or Word files.
It can be used as a standalone utility to extract tender data for the tender reminder feature.

Usage:
    python tender_file_parser.py <file_path>

Supported formats:
    - CSV (with columns: tender_name, email, bidding_date)
    - Excel (.xlsx, .xls)
    - Word (.docx) [if python-docx is installed]

Example:
    python tender_file_parser.py test_files/sample_tenders.csv
"""
import sys
import os
import re
import pandas as pd
from datetime import datetime

try:
    from docx import Document
    DOCX_SUPPORTED = True
except ImportError:
    DOCX_SUPPORTED = False

REQUIRED_COLUMNS = {'tender_name', 'email', 'bidding_date'}
EMAIL_REGEX = r"^[\w\.-]+@[\w\.-]+\.\w+$"

def format_error(msg):
    return f"\u274C {msg}"

def format_warning(msg):
    return f"\u26A0\uFE0F {msg}"

def format_info(msg):
    return f"\u2139\uFE0F {msg}"

def validate_columns(df):
    missing = REQUIRED_COLUMNS - set(df.columns)
    if missing:
        raise ValueError(format_error(f"Missing required columns: {', '.join(missing)}. Please ensure your file includes all required columns: tender_name, email, and bidding_date."))

def validate_email(email):
    return re.match(EMAIL_REGEX, str(email)) is not None

def validate_date(date_str):
    try:
        # Try parsing common date formats
        datetime.strptime(str(date_str), "%d/%m/%Y")
        return True
    except Exception:
        try:
            datetime.strptime(str(date_str), "%Y-%m-%d")
            return True
        except Exception:
            return False

def parse_csv(file_path):
    df = pd.read_csv(file_path)
    if df.empty:
        raise ValueError(format_error("The file is empty. Please upload a file with at least one tender entry."))
    validate_columns(df)
    tenders = []
    seen = set()
    warnings = []
    for idx, row in df.iterrows():
        tender = row.to_dict()
        errors = []
        for col in REQUIRED_COLUMNS:
            if not str(tender.get(col, '')).strip():
                errors.append(f"{col} is empty")
        if not validate_email(tender.get('email', '')):
            errors.append(f"Invalid email address '{tender.get('email', '')}'")
        if not validate_date(tender.get('bidding_date', '')):
            errors.append(f"Invalid bidding date '{tender.get('bidding_date', '')}' (use DD/MM/YYYY)")
        key = (tender.get('tender_name', '').strip().lower(), tender.get('email', '').strip().lower())
        if key in seen:
            errors.append('Duplicate tender (same name and email)')
        else:
            seen.add(key)
        try:
            bid_date = datetime.strptime(str(tender.get('bidding_date', '')), "%d/%m/%Y")
        except Exception:
            try:
                bid_date = datetime.strptime(str(tender.get('bidding_date', '')), "%Y-%m-%d")
            except Exception:
                bid_date = None
        if bid_date and bid_date.date() < datetime.now().date():
            errors.append(f"Bidding date '{tender.get('bidding_date', '')}' is in the past")
        if errors:
            warnings.append(format_warning(f"Row {idx+1} skipped: {', '.join(errors)}"))
            continue
        tenders.append(tender)
    for w in warnings:
        print(w)
    if not tenders:
        raise ValueError(format_error("No valid tenders found in the file after validation. Please review your data for errors and try again."))
    if warnings:
        print(format_info(f"Summary: {len(warnings)} row(s) skipped, {len(tenders)} row(s) parsed successfully."))
    return tenders

def parse_excel(file_path):
    df = pd.read_excel(file_path)
    if df.empty:
        raise ValueError(format_error("The file is empty. Please upload a file with at least one tender entry."))
    validate_columns(df)
    tenders = []
    seen = set()
    warnings = []
    for idx, row in df.iterrows():
        tender = row.to_dict()
        errors = []
        for col in REQUIRED_COLUMNS:
            if not str(tender.get(col, '')).strip():
                errors.append(f"{col} is empty")
        if not validate_email(tender.get('email', '')):
            errors.append(f"Invalid email address '{tender.get('email', '')}'")
        if not validate_date(tender.get('bidding_date', '')):
            errors.append(f"Invalid bidding date '{tender.get('bidding_date', '')}' (use DD/MM/YYYY)")
        key = (tender.get('tender_name', '').strip().lower(), tender.get('email', '').strip().lower())
        if key in seen:
            errors.append('Duplicate tender (same name and email)')
        else:
            seen.add(key)
        try:
            bid_date = datetime.strptime(str(tender.get('bidding_date', '')), "%d/%m/%Y")
        except Exception:
            try:
                bid_date = datetime.strptime(str(tender.get('bidding_date', '')), "%Y-%m-%d")
            except Exception:
                bid_date = None
        if bid_date and bid_date.date() < datetime.now().date():
            errors.append(f"Bidding date '{tender.get('bidding_date', '')}' is in the past")
        if errors:
            warnings.append(format_warning(f"Row {idx+1} skipped: {', '.join(errors)}"))
            continue
        tenders.append(tender)
    for w in warnings:
        print(w)
    if not tenders:
        raise ValueError(format_error("No valid tenders found in the file after validation. Please review your data for errors and try again."))
    if warnings:
        print(format_info(f"Summary: {len(warnings)} row(s) skipped, {len(tenders)} row(s) parsed successfully."))
    return tenders

def parse_docx(file_path):
    if not DOCX_SUPPORTED:
        raise ImportError(format_error("python-docx is not installed. Please install it to parse Word files."))
    doc = Document(file_path)
    tenders = []
    headers = []
    seen = set()
    warnings = []
    table_count = len(doc.tables)
    if table_count == 0:
        raise ValueError(format_error("No tables found in Word file."))
    for i, table in enumerate(doc.tables):
        if len(table.rows) == 0:
            warnings.append(format_warning(f"Table {i+1} is empty."))
            continue
        for j, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            if j == 0:
                headers = cells
                headers = [h.lower().strip() for h in headers]
            else:
                if len(cells) != len(headers):
                    warnings.append(format_warning(f"Row {j+1} in table {i+1} skipped: Number of columns does not match the header."))
                    continue
                tender = dict(zip(headers, cells))
                tender = {k.lower().strip(): v for k, v in tender.items()}
                if not REQUIRED_COLUMNS.issubset(tender.keys()):
                    warnings.append(format_warning(f"Row {j+1} in table {i+1} skipped: Missing required columns."))
                    continue
                errors = []
                for col in REQUIRED_COLUMNS:
                    if not str(tender.get(col, '')).strip():
                        errors.append(f"{col} is empty")
                if not validate_email(tender.get('email', '')):
                    errors.append(f"Invalid email address '{tender.get('email', '')}'")
                if not validate_date(tender.get('bidding_date', '')):
                    errors.append(f"Invalid bidding date '{tender.get('bidding_date', '')}' (use DD/MM/YYYY)")
                key = (tender.get('tender_name', '').strip().lower(), tender.get('email', '').strip().lower())
                if key in seen:
                    errors.append('Duplicate tender (same name and email)')
                else:
                    seen.add(key)
                try:
                    bid_date = datetime.strptime(str(tender.get('bidding_date', '')), "%d/%m/%Y")
                except Exception:
                    try:
                        bid_date = datetime.strptime(str(tender.get('bidding_date', '')), "%Y-%m-%d")
                    except Exception:
                        bid_date = None
                if bid_date and bid_date.date() < datetime.now().date():
                    errors.append(f"Bidding date '{tender.get('bidding_date', '')}' is in the past")
                if errors:
                    warnings.append(format_warning(f"Row {j+1} in table {i+1} skipped: {', '.join(errors)}"))
                    continue
                tenders.append(tender)
    for w in warnings:
        print(w)
    if not tenders:
        raise ValueError(format_error("No valid tenders found in the file after validation. Please review your data for errors and try again."))
    if warnings:
        print(format_info(f"Summary: {len(warnings)} row(s) skipped, {len(tenders)} row(s) parsed successfully."))
    return tenders

def parse_tender_file(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    if ext == '.csv':
        return parse_csv(file_path)
    elif ext in ['.xlsx', '.xls']:
        return parse_excel(file_path)
    elif ext == '.docx':
        return parse_docx(file_path)
    else:
        raise ValueError(format_error(f"Unsupported file type: '{ext}'. Supported types: .csv, .xlsx, .xls, .docx."))

def main():
    if len(sys.argv) < 2:
        print(format_info("Usage: python tender_file_parser.py <file_path>"))
        sys.exit(1)
    file_path = sys.argv[1]
    if not os.path.exists(file_path):
        print(format_error(f"File not found: '{file_path}'. Please check the file path and try again."))
        sys.exit(1)
    try:
        tenders = parse_tender_file(file_path)
        print(format_info(f"Parsed {len(tenders)} tenders:"))
        for tender in tenders:
            print(tender)
    except Exception as e:
        import traceback
        print(format_error(f"An error occurred: {e}"))
        print("\u26A0\uFE0F Please check your file for corruption or formatting issues. (See details below for troubleshooting.)")
        # traceback.print_exc()
        sys.exit(1)

if __name__ == "__main__":
    main()
